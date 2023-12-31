from io import StringIO
import re
import streamlit as st
from PyPDF2 import PdfReader
import requests
from languages import languages
from docx import Document
from gtts import gTTS

key = st.secrets['LANGUAGE_KEY']
endpoint = st.secrets['LANGUAGE_ENDPOINT']

from azure.ai.textanalytics import TextAnalyticsClient
from azure.core.credentials import AzureKeyCredential

st.set_page_config( 
     page_title="EduVerse", 
     page_icon="🏫",
     layout="wide", 
     initial_sidebar_state="expanded", 
 ) 

# Authenticate the client using your key and endpoint 
def authenticate_client():
    ta_credential = AzureKeyCredential(key)
    text_analytics_client = TextAnalyticsClient(
            endpoint=endpoint, 
            credential=ta_credential)
    return text_analytics_client

client = authenticate_client()

if 'username' not in st.session_state:
    st.session_state["username"]="User"

hide_default_format = """
       <style>
       #MainMenu {visibility: hidden; }
       footer {visibility: hidden;}
       </style>
       """
st.markdown(hide_default_format, unsafe_allow_html=True)

def create_sublists(paragraph, sublist_size, sub_size=3):
    sentences = paragraph.split('. ')  # Split at period followed by a space
    sentences = [s.strip() for s in sentences]  # Remove leading/trailing spaces

    # Split further at exclamation marks and question marks
    split_sentences = []
    for sentence in sentences:
        split_sentences.extend(sentence.split('! '))
        split_sentences.extend(sentence.split('? '))

    split_sentences = [s.strip() for s in split_sentences]  # Remove leading/trailing spaces
    
    original_list = split_sentences
    sublists = []
    sublist = []

    s=''
    c=0
    for element in original_list:
        s=s+element
        c+=1
        if c==sub_size:
            sublist.append(s)
            s=''
            c=0
        if len(sublist) == sublist_size:
            sublists.append(sublist)
            sublist = []
    
    # Add the remaining elements if the original list length is not divisible by sublist_size
    if s!='':
        sublist.append(s)
        sublists.append(sublist)
    if len(sublist) > 0:
        sublists.append(sublist)

    return sublists

def summarizer(text):
    # [START extract_summary]
    import os
    from azure.core.credentials import AzureKeyCredential
    from azure.ai.textanalytics import TextAnalyticsClient

    text_analytics_client = TextAnalyticsClient(
        endpoint=endpoint,
        credential=AzureKeyCredential(key),
    )

    chunks = create_sublists(text, 7,3)
    summaries = []
    for chunk in chunks:
        poller = text_analytics_client.begin_extract_summary(chunk)
        extract_summary_results = poller.result()
        for result in extract_summary_results:
            if result.kind == "ExtractiveSummarization":
                return(" ".join([sentence.text for sentence in result.sentences]))
    
    return ' '.join(summaries)

def translator(text, to_lang):
    import uuid
    path = '/translate?api-version=3.0'
    params = '&to='+to_lang
    endpoint='https://api.cognitive.microsofttranslator.com/'
    constructed_url = endpoint + path + params

    headers = {
        'Ocp-Apim-Subscription-Key': st.secrets['TRANSLATOR_KEY'],
        'Ocp-Apim-Subscription-Region': "centralindia",
        'Content-type': 'application/json',
        'X-ClientTraceId': str(uuid.uuid4())
    }

    # You can pass more than one object in body.
    body = [{
        'text' : text
    }]
    request = requests.post(constructed_url, headers=headers, json=body)
    response = request.json()

    return(response[0]['translations'][0]['text'])

def main():

    def gradient_text(text, color1, color2):
        gradient_css = f"""
        background: -webkit-linear-gradient(left, {color1}, {color2});
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
        font-weight: bold;
        font-size: 42px;
        """
        return f'<span style="{gradient_css}">{text}</span>'
    
    def gradient(text, color1, color2):
        gradient_css = f"""
        background: -webkit-linear-gradient(left, {color2}, {color1});
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
        font-size: 22px;
        """
        return f'<span style="{gradient_css}">{text}</span>'


    color1 = "#0d3270"
    color2 = "#0fab7b"
    text = "LinguaBrief"

    styled_text = gradient_text(text, color1, color2)
    st.write(f"<div style='text-align: center;'>{styled_text}</div>", unsafe_allow_html=True)
    styled_text = gradient("Get your Text Summarized and Translated", color1, color2)
    st.write(f"<div style='text-align: center;'>{styled_text}</div>", unsafe_allow_html=True)
    st.subheader("Weclome "+st.session_state["username"]+" 👋")

    Option = st.selectbox("Choose the format", ('Upload PDF','Text Input'), index=1)

    if Option == 'Text Input':
        txt = st.text_area('Text to summarize', '')
        if 'selected_language' not in st.session_state:
            st.session_state.selected_language = 'English'
        to_lang = st.session_state.selected_language
        to_lang = st.selectbox("Select the language", languages.values(), index=list(languages.values()).index(st.session_state.selected_language))
        col1,  col2, col3 = st.columns(3)
        if col1.button('SUMMARIZE'):
            if txt is None:
                st.error("ERROR: ENTER SOME TEXT", icon="🚨")
                st.stop()
            if txt is not None:
                summarized_text = summarizer(txt)
                summarized_text = translator(summarized_text, to_lang=list(languages.keys())[list(languages.values()).index(to_lang)])
                if summarized_text:
                    st.success(" Successfully summarized!", icon='✅')
                    st.header('Result')
                    st.write(summarized_text)
                    doc = Document()
                    doc.add_heading('Here\'s The Note ', level=1)
                    doc.add_paragraph(summarized_text)
                    doc.save('note.doc')
                    tts = gTTS(text=summarized_text, lang='en')
                    tts.save("outputaudio.mp3")
                    audio_file = open('outputaudio.mp3', 'rb')
                    audio_bytes = audio_file.read()
                    st.audio(audio_bytes, format='audio/,mp3')
                    st.download_button(
                    label="Download Note",
                    data=open("note.doc", "rb").read(),
                    file_name="note.doc",
                    mime="application/octet-stream",
                    help="Click to download the summarized and translated note."
                    )   

                    # Download mp3 file
                    with open('outputaudio.mp3', 'rb') as file:
                        output = file.read()
        
                    st.download_button(
                        label="Download Audio",
                        data=output,
                        file_name='output.mp3',
                        mime="audio/mpeg",
                        help="To download an audio file of the generated results"
                    )


    elif Option == 'Upload PDF':
        file = st.file_uploader("Choose a file")
        if 'selected_language' not in st.session_state:
            st.session_state.selected_language = 'English'
        to_lang = st.session_state.selected_language
        to_lang = st.selectbox("Select the language for summary", languages.values(), index=list(languages.values()).index(st.session_state.selected_language))
        if file:
            if file.name[-3:] == "pdf":
                pdfReader = PdfReader(file)
                num = len(pdfReader.pages)
                text=''
                for i in range(0,num):
                    pageobj = pdfReader.pages[i]
                    resulttext = pageobj.extract_text()
                    text = text.join(resulttext)
                    text = re.sub(r'(?<=\S)\s{2,}(?=\S)', ' ', text)
                    text = re.sub(r'\n', ' ', text)
                    summarized_text = summarizer(text)
                      
            else:
                stringio = StringIO(file.getvalue().decode("utf-8"))
                text=stringio.read()
                summarized_text = summarizer(text)
                    
            if summarized_text:
                summarized_text = translator(summarized_text, to_lang=list(languages.keys())[list(languages.values()).index(to_lang)])
                st.success(" Successfully summarized!", icon='✅')
                st.header('Result')
                st.write(summarized_text)
                doc = Document()
                doc.add_heading('Here\'s The Note ', level=1)
                doc.add_paragraph(summarized_text)
                doc.save('note.doc')
                tts = gTTS(text=summarized_text, lang='en')
                tts.save("outputaudio.mp3")
                audio_file = open('outputaudio.mp3', 'rb')
                audio_bytes = audio_file.read()
                st.audio(audio_bytes, format='audio/,mp3')
                st.download_button(
                label="Download Note",
                data=open("note.doc", "rb").read(),
                file_name="note.doc",
                mime="application/octet-stream",
                help="Click to download the summarized and translated note."
                )   

                # Download mp3 file
                with open('outputaudio.mp3', 'rb') as file:
                    output = file.read()
        
                st.download_button(
                    label="Download Audio",
                    data=output,
                    file_name='output.mp3',
                    mime="audio/mpeg",
                    help="To download an audio file of the generated results"
                )


if __name__ == "__main__":
    main()
footer="""<style>

a:hover,  a:active {
color: red;
background-color: transparent;
text-decoration: underline;
}

.footer {
position: fixed;
left: 0;
bottom: 0;
width: 100%;
background-color: white;
color: black;
text-align: center;
}
</style>
<div class="footer">
<p>Developed with ❤️ by <a style='display: inline; text-align: center;' href="https://www.linkedin.com/in/harshavardhan-bajoria/" target="_blank">Harshavardhan Bajoria</a></p>
</div>
"""
st.markdown(footer,unsafe_allow_html=True)
